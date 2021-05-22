# GUIWiki.py
import wikipedia

# python to docx
from docx import Document
def Wiki(keyword,lang='lo'):
    wikipedia.set_lang(lang)
    # summary ສຳຫຼັບບົດຄວາມສະຫຼຸບ
    data = wikipedia.summary(keyword)

    # page + content ບົດຄວາມທັ້ງໜ້າ
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #ສ້າງໄຟລ໌ word ໃນ python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('ບັນທືກໄຟລ໌ສຳເລັດ')


# ປ່ຽນເປັນພາສາລາວ
wikipedia.set_lang('lo')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('ໂປຣເເກຣມຄົ້ນຫາບົດຄວາມ')
GUI.geometry('400x300')

FONT1 = ('Phetsarath OT',15)


L = ttk.Label(GUI, text='ຄົ້ນຫາບົດຄວາມ',font=FONT1)
L.pack()
# ຊ່ອງຄົ້ນຫາຂໍ້ມູນ
v_search = StringVar() 
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=40)
E1.pack(pady=10)

# ปุ่มค้นหา
def Search():
    keyword = v_search.get() # .get() ຄືດຶງຂໍ້ມູນເຂົ້າມາ 
    try:
        # ລອງຄົ້ນຫາເບິ່ງວ່າ ໄດ້ຜົນລັບ ຫຼືບໍ ຫາກໄດ້ໃຫ້ຜ່ານໄປ 
        language = v_radio.get() # lo / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('ບັນທືກໄຟລ໌ສຳເລັດ','ຄົ້ນຫາຂໍ້ມູນສຳເລັດ ບັນທຶກຮຽບຮ້ອຍເເລ້ວ')
    except:
        # ຫາກຣັນຄຳສັ່ງເເລ້ວມີບັນຫາ ສະເເດງ ຂໍ້ຄວາມເເຈ້ງຕືນ 
        messagebox.showwarning('Keyword Error','ກະລຸນາກອກຄຳຄົ້ນຫາໃຫມ່')
        
    # print(wikipedia.search(keyword))
    # result = wikipedia.summary(keyword)
    # print(result)

B1 = ttk.Button(GUI,text='Search',command=Search)
B1.pack(ipadx=20,ipady=10) 

# ເລືອກພາສາ
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar() 

RB1 = ttk.Radiobutton(F1,text='ລາວ',variable=v_radio,value='lo')
RB2 = ttk.Radiobutton(F1,text='ອັງກິດ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='ຈີນ',variable=v_radio,value='zh')
RB1.invoke() # ສັ່ງຄ່າເລີມເປັນພາສາລາວ

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)


GUI.mainloop()